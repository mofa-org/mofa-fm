"""
文件解析服务 - 从各种格式文件中提取文本内容
支持: txt, pdf, md, docx
"""
import os
from typing import Dict


class FileParser:
    """文件解析器"""

    @staticmethod
    def parse(file_path: str, file_type: str = None) -> Dict[str, any]:
        """
        解析文件并提取文本内容

        Args:
            file_path: 文件路径
            file_type: 文件类型 (txt/pdf/md/docx)，如果为None则自动检测

        Returns:
            {
                'success': bool,
                'text': str,  # 提取的文本内容
                'error': str  # 错误信息（如果失败）
            }
        """
        if not os.path.exists(file_path):
            return {'success': False, 'error': '文件不存在'}

        # 自动检测文件类型
        if not file_type:
            ext = os.path.splitext(file_path)[1].lower()
            file_type = ext.lstrip('.')

        # 根据类型解析
        try:
            if file_type in ['txt', 'md']:
                text = FileParser._parse_text(file_path)
            elif file_type == 'pdf':
                text = FileParser._parse_pdf(file_path)
            elif file_type in ['docx', 'doc']:
                text = FileParser._parse_docx(file_path)
            else:
                return {'success': False, 'error': f'不支持的文件类型: {file_type}'}

            return {
                'success': True,
                'text': text,
                'char_count': len(text)
            }

        except Exception as e:
            return {'success': False, 'error': f'解析失败: {str(e)}'}

    @staticmethod
    def _parse_text(file_path: str) -> str:
        """解析文本文件 (txt, md)"""
        encodings = ['utf-8', 'gbk', 'gb2312', 'utf-16']

        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue

        # 如果所有编码都失败，使用二进制模式读取
        with open(file_path, 'rb') as f:
            content = f.read()
            # 尝试解码，忽略错误
            return content.decode('utf-8', errors='ignore')

    @staticmethod
    def _parse_pdf(file_path: str) -> str:
        """解析PDF文件"""
        try:
            import PyPDF2

            text_parts = []
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text_parts.append(page.extract_text())

            return '\n\n'.join(text_parts)

        except ImportError:
            raise Exception('缺少 PyPDF2 库，请运行: pip install PyPDF2')

    @staticmethod
    def _parse_docx(file_path: str) -> str:
        """解析Word文档 (docx)"""
        try:
            import docx

            doc = docx.Document(file_path)
            text_parts = []

            # 提取段落
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # 提取表格
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)

            return '\n\n'.join(text_parts)

        except ImportError:
            raise Exception('缺少 python-docx 库，请运行: pip install python-docx')

    @staticmethod
    def get_file_info(file_path: str) -> Dict:
        """
        获取文件基本信息

        Returns:
            {
                'filename': str,
                'size': int,  # 字节
                'extension': str
            }
        """
        return {
            'filename': os.path.basename(file_path),
            'size': os.path.getsize(file_path),
            'extension': os.path.splitext(file_path)[1].lstrip('.')
        }


# 测试代码
if __name__ == "__main__":
    # 测试文本文件
    test_content = "这是一个测试文件\n包含多行文本"
    with open('/tmp/test.txt', 'w', encoding='utf-8') as f:
        f.write(test_content)

    parser = FileParser()
    result = parser.parse('/tmp/test.txt')
    print("解析结果:", result)
